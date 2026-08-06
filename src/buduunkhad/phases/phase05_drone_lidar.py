"""Phase 05 — acquired drone/LiDAR intake, terrain QA and field handoff.

Physical acquisition is outside the pipeline.  This phase starts from the delivered DJI Terra,
Metashape, control and daily-report tree and produces compact products that can be sealed,
reviewed and delivered without copying the campaign's raw photos or LAS files.
"""

from __future__ import annotations

import json
import os
from pathlib import Path

from buduunkhad.core import (
    pdf_map,
    phase05_drone,
    qgis_project,
    raster_writers,
    registers,
    vector_io,
)
from buduunkhad.core.qaqc import Decision, QAQCReport, new_report
from buduunkhad.phases.base import Phase, PhaseResult, RunContext

_BLOCK_COLUMNS = [
    "block_id",
    "block_number",
    "sensor",
    "dem_path",
    "dsm_path",
    "dom_path",
    "preferred_ortho_path",
    "colour_corrected_ortho",
    "point_cloud_path",
    "point_count",
    "point_cloud_crs",
    "quality_report_path",
    "core_products_complete",
    "raster_failures",
    "crs_mismatches",
    "data_gaps",
    "status",
]


class Phase05DroneLidar(Phase):
    id = "05"
    name = "DJI Matrice 400 Drone / LiDAR / Photogrammetry Survey"
    mode = "orchestrate"
    input_numbers = [*range(9, 23), *range(75, 79)]
    custom_subfolders = [
        "01_Source_Inventory",
        "02_Flight_and_Control_Records",
        "03_Raster_Products",
        "04_Point_Cloud_Products",
        "05_GIS_and_3D",
        "06_Field_Planning",
        "07_Final_Deliverables",
        "08_QAQC_and_Confidence",
    ]
    gate_condition = (
        "Acquired survey products inventoried, spatial products verified in EPSG:32647, "
        "and usable terrain/GIS handoff produced."
    )

    def __init__(self) -> None:
        self._survey: phase05_drone.DroneSurvey | None = None
        self._block_rows: list[dict[str, object]] = []
        self._terrain_outputs: tuple[Path, ...] = ()
        self._coverage_path: Path | None = None
        self._interpretation_path: Path | None = None
        self._project_path: Path | None = None
        self._prospect_count = 0

    def run(self, ctx: RunContext) -> PhaseResult:
        self.__init__()
        result = PhaseResult(self.id, status="dry-run" if ctx.dry_run else "ok")
        root = ctx.phase_dir(self.id)
        register_path = (
            root
            / "01_Source_Inventory"
            / f"{ctx.config.register_prefix}_Drone_Block_Asset_Register.xlsx"
        )
        if ctx.dry_run:
            registers.write_table_xlsx(
                [], _BLOCK_COLUMNS, register_path, sheet_title="Drone blocks"
            )
            plan = (
                root
                / "07_Final_Deliverables"
                / (f"{ctx.config.register_prefix}_Drone_Flight_Plan.pdf")
            )
            pdf_map.write_text_pdf(
                plan,
                [
                    "Phase 05 acquired drone/LiDAR campaign - dry run",
                    "No survey source was opened.",
                    f"Target CRS: {ctx.config.crs.target_authority}",
                ],
            )
            report = self._write_handoff_report(ctx, root, dry_run=True)
            result.outputs.extend((register_path, plan, report))
            result.log("dry-run Phase 05 paths and register schema created")
            return result

        source_root = ctx.config.phase05_source_root
        if source_root is None:
            raise phase05_drone.Phase05DataError(
                "Phase 05 requires BUDUUNKHAD_PHASE05_SOURCE_ROOT pointing to the acquired "
                "Drone campaign directory"
            )
        phase05_drone.validate_source_separation(
            source_root,
            ctx.config.output_root,
            ctx.config.runs_root,
            ctx.config.results_root,
        )
        survey = phase05_drone.inspect_survey(
            source_root,
            target_epsg=ctx.config.target_epsg,
        )
        self._survey = survey
        self._block_rows = phase05_drone.block_rows(
            survey,
            target_epsg=ctx.config.target_epsg,
        )
        registers.write_table_xlsx(
            self._block_rows,
            _BLOCK_COLUMNS,
            register_path,
            sheet_title="Drone blocks",
            widths=[12, 10, 8, 45, 45, 45, 45, 14, 45, 16, 18, 45, 16, 35, 35, 35, 12],
        )
        source_record = (
            root
            / "01_Source_Inventory"
            / (f"{ctx.config.register_prefix}_Drone_Source_Inventory.json")
        )
        source_record.write_text(
            json.dumps(
                {
                    "format_version": "1.0.0",
                    "project": ctx.config.project.name,
                    "run_id": ctx.run_id,
                    "source_root": str(survey.layout.root),
                    "source_tree_identity": phase05_drone.source_tree_identity(survey.layout.root),
                    "identity_basis": "relative path, byte size and filesystem modification time",
                    "block_count": len(survey.blocks),
                    "control_file_count": len(survey.control_files),
                    "daily_report_file_count": len(survey.daily_report_files),
                    "final_dem": survey.final_dem.relative_path,
                    "final_dem_size_bytes": survey.final_dem.size_bytes,
                    "limitations": [
                        "Large campaign assets are read in place and are not duplicated in the run.",
                        "The metadata identity avoids a full 624 GB content hash; exported derivatives are fully sealed.",
                    ],
                },
                indent=2,
                sort_keys=True,
            )
            + "\n",
            encoding="utf-8",
        )

        coverage = phase05_drone.coverage_geodataframe(
            survey,
            target_epsg=ctx.config.target_epsg,
        )
        if coverage.empty:
            raise phase05_drone.Phase05DataError("no valid block coverage could be derived")
        self._coverage_path = (
            root
            / "05_GIS_and_3D"
            / ctx.data_name(
                "Drone_Block_Coverage",
                crs_or_param=f"EPSG{ctx.config.target_epsg}",
                ext="gpkg",
            )
        )
        vector_io.write_layer(coverage, self._coverage_path, layer="drone_block_coverage")

        raster_dir = root / "03_Raster_Products"
        dtm = raster_dir / ctx.data_name(
            "Drone_DTM",
            crs_or_param=f"EPSG{ctx.config.target_epsg}",
            ext="tif",
        )
        hillshade = raster_dir / ctx.data_name(
            "Drone_DTM_Hillshade",
            crs_or_param=f"EPSG{ctx.config.target_epsg}",
            ext="tif",
        )
        slope = raster_dir / ctx.data_name(
            "Drone_DTM_SlopeDeg",
            crs_or_param=f"EPSG{ctx.config.target_epsg}",
            ext="tif",
        )
        self._terrain_outputs = phase05_drone.write_terrain_products(
            survey.final_dem.path,
            dtm_output=dtm,
            hillshade_output=hillshade,
            slope_output=slope,
        )

        self._interpretation_path = self._write_interpretation_package(ctx, root, coverage)
        flight_plan = self._write_coverage_map(ctx, root, coverage)
        self._project_path = self._write_qgis_project(
            ctx,
            root,
            dtm=dtm,
            hillshade=hillshade,
            slope=slope,
            coverage_path=self._coverage_path,
            interpretation_path=self._interpretation_path,
            extent=survey.final_dem.bounds,
        )
        handoff = self._write_handoff_report(ctx, root, dry_run=False)

        result.outputs.extend(
            (
                register_path,
                source_record,
                self._coverage_path,
                *self._terrain_outputs,
                self._interpretation_path,
                flight_plan,
                self._project_path,
                handoff,
            )
        )
        missing_clouds = sum(block.point_cloud is None for block in survey.blocks)
        result.log(
            f"inventoried {len(survey.blocks)} blocks; {missing_clouds} missing point cloud(s); "
            f"{len(survey.daily_report_files)} daily report file(s)"
        )
        return result

    def qaqc(self, ctx: RunContext) -> QAQCReport:
        report = new_report(self.id, self.name)
        if ctx.dry_run:
            report.add(
                "Phase 05 output layout",
                "Dry-run register, coverage-plan path and handoff path created.",
                decision=Decision.PASS,
            )
            return report
        assert self._survey is not None
        report.add(
            "Acquired block inventory",
            "Every B-numbered processed block is represented once in the block register.",
            decision=Decision.PASS if len(self._survey.blocks) == 79 else Decision.FAIL,
            note=f"{len(self._survey.blocks)} block(s) found.",
        )
        invalid_rasters = sum(
            bool(row["raster_failures"] or row["crs_mismatches"]) for row in self._block_rows
        )
        report.add(
            "Raster open and CRS checks",
            f"DEM, DSM, DOM and selected orthos open in EPSG:{ctx.config.target_epsg}.",
            decision=Decision.PASS if invalid_rasters == 0 else Decision.FAIL,
            note=f"{invalid_rasters} block(s) contain a raster failure or CRS mismatch.",
        )
        report.add(
            "Terrain handoff",
            "DTM, hillshade and slope are valid sealed Cloud-Optimized GeoTIFFs.",
            decision=(
                Decision.PASS
                if self._terrain_outputs
                and all(raster_writers.is_cog(path) for path in self._terrain_outputs)
                else Decision.FAIL
            ),
            note=f"{len(self._terrain_outputs)} terrain raster(s).",
        )
        point_cloud_count = sum(block.point_cloud is not None for block in self._survey.blocks)
        invalid_clouds = sum(
            block.point_cloud is not None and block.point_cloud.open_status == "fail"
            for block in self._survey.blocks
        )
        report.add(
            "Point-cloud source readiness",
            "Available LAS/LAZ headers open and their gaps remain explicit.",
            decision=Decision.PASS if invalid_clouds == 0 else Decision.FAIL,
            note=(
                f"{point_cloud_count}/{len(self._survey.blocks)} block point clouds; "
                "LAZ/COPC conversion is not part of this first processing slice."
            ),
        )
        report.add(
            "GIS and 3D handoff",
            "Coverage, interpretation package, PDF map and QGIS terrain project were produced.",
            decision=(
                Decision.PASS
                if all(
                    path is not None and path.is_file()
                    for path in (
                        self._coverage_path,
                        self._interpretation_path,
                        self._project_path,
                    )
                )
                else Decision.FAIL
            ),
            note=f"{self._prospect_count} Phase 04 prospect polygon(s) embedded.",
        )
        report.add(
            "Full-resolution orthomosaic and DSM mosaics",
            "Master deliverables are generated from the acquired per-block sources.",
            decision=Decision.NA,
            note="Scheduled for the next Phase 05 processing slice; per-block sources are inventoried.",
        )
        return report

    def _write_interpretation_package(self, ctx: RunContext, root: Path, coverage) -> Path:
        import fiona
        from pyproj import CRS

        path = (
            root
            / "05_GIS_and_3D"
            / ctx.data_name(
                "Drone_Structure_Outcrop_Interpretation",
                crs_or_param=f"EPSG{ctx.config.target_epsg}",
                ext="gpkg",
            )
        )
        vector_io.write_layer(coverage, path, layer="drone_block_coverage")
        generic_props = {
            "feature_id": "str:64",
            "source": "str:254",
            "confidence": "str:32",
            "notes": "str:254",
        }
        crs_wkt = CRS.from_epsg(ctx.config.target_epsg).to_wkt()
        for name, geometry in (
            ("structure_interpretation_line", "LineString"),
            ("outcrop_interpretation_polygon", "MultiPolygon"),
            ("field_observation_point", "Point"),
            ("sample_planning_point", "Point"),
        ):
            with fiona.open(
                path,
                "w",
                driver="GPKG",
                layer=name,
                schema={"geometry": geometry, "properties": generic_props},
                crs_wkt=crs_wkt,
            ):
                pass

        prospect_source = self._phase04_prospects(ctx)
        if prospect_source is not None:
            prospects = vector_io.read_layer(prospect_source, "prospect_candidate_areas")
            if prospects.crs is None:
                raise phase05_drone.Phase05DataError("Phase 04 prospect layer has no CRS")
            prospects = prospects.to_crs(epsg=ctx.config.target_epsg)
            if not prospects.empty:
                vector_io.write_layer(
                    prospects,
                    path,
                    layer="phase04_prospect_candidates",
                    mode="w",
                )
                self._prospect_count = len(prospects)
        return path

    def _phase04_prospects(self, ctx: RunContext) -> Path | None:
        phase04 = ctx.phase_dir("04")
        for path in sorted(phase04.rglob("*.gpkg")):
            try:
                if "prospect_candidate_areas" in vector_io.list_gpkg_layers(path):
                    return path
            except Exception:  # noqa: BLE001 - unrelated GeoPackages are not prospect inputs
                continue
        return None

    def _write_coverage_map(self, ctx: RunContext, root: Path, coverage) -> Path:
        path = (
            root / "07_Final_Deliverables" / (f"{ctx.config.register_prefix}_Drone_Flight_Plan.pdf")
        )
        labels = []
        for row in coverage.itertuples():
            point = row.geometry.representative_point()
            labels.append(pdf_map.MapLabel(float(point.x), float(point.y), str(row.block_id)))
        return pdf_map.write_map_pdf(
            path,
            title="Phase 5 Acquired Drone Coverage and Flight Record",
            subtitle=f"{ctx.config.project.project_code} / {ctx.config.project.name}",
            crs_label=ctx.config.crs.target_authority,
            run_id=ctx.run_id,
            layers=(
                pdf_map.MapLayer(
                    name="Acquired drone blocks",
                    geometries=tuple(coverage.geometry),
                    stroke=(0.1, 0.3, 0.8),
                    fill=(0.75, 0.85, 1.0),
                    line_width=0.8,
                ),
            ),
            labels=tuple(labels),
            notes=(
                "Retrospective coverage derived from the acquired orthomosaic footprints.",
                "This map records completed coverage; it does not authorize a future flight.",
            ),
            footer="Acquired campaign record and field-planning support product.",
        )

    def _write_qgis_project(
        self,
        ctx: RunContext,
        root: Path,
        *,
        dtm: Path,
        hillshade: Path,
        slope: Path,
        coverage_path: Path,
        interpretation_path: Path,
        extent: tuple[float, float, float, float],
    ) -> Path:
        path = (
            root
            / "07_Final_Deliverables"
            / (f"{ctx.config.register_prefix}_Drone_3D_Field_Project.qgz")
        )

        def source(value: Path, layer: str | None = None) -> str:
            relative = os.path.relpath(value, path.parent).replace("\\", "/")
            return f"{relative}|layername={layer}" if layer else relative

        layers = [
            qgis_project.QgzLayer(
                "Drone DTM",
                source(dtm),
                "Raster",
                visible=False,
                group="Terrain",
                provider="gdal",
            ),
            qgis_project.QgzLayer(
                "Drone hillshade",
                source(hillshade),
                "Raster",
                group="Terrain",
                provider="gdal",
            ),
            qgis_project.QgzLayer(
                "Drone slope",
                source(slope),
                "Raster",
                visible=False,
                group="Terrain",
                provider="gdal",
            ),
            qgis_project.QgzLayer(
                "Drone block coverage",
                source(coverage_path, "drone_block_coverage"),
                "MultiPolygon",
                symbol=qgis_project.polygon_outline("20,80,220,255", 0.6),
                group="Survey coverage",
            ),
        ]
        if self._prospect_count:
            layers.append(
                qgis_project.QgzLayer(
                    "Phase 04 prospect candidates",
                    source(interpretation_path, "phase04_prospect_candidates"),
                    "MultiPolygon",
                    symbol=qgis_project.polygon_outline("220,60,20,255", 0.8),
                    group="Field planning",
                )
            )
        return qgis_project.write_layered_qgz(
            path,
            epsg=ctx.config.target_epsg,
            title=f"{ctx.config.project.name} Phase 05 Drone 3D and Field Planning",
            layers=layers,
            initial_extent=qgis_project.QgzExtent(*extent),
            terrain_layer_name="Drone DTM",
            qgis_version="4.0.3",
        )

    def _write_handoff_report(self, ctx: RunContext, root: Path, *, dry_run: bool) -> Path:
        from docx import Document

        path = (
            root
            / "07_Final_Deliverables"
            / (f"{ctx.config.register_prefix}_Drone_Survey_Handoff_Report.docx")
        )
        doc = Document()
        doc.add_heading("Phase 05 Drone / LiDAR Survey Handoff", level=1)
        doc.add_paragraph(f"Project: {ctx.config.project.name} ({ctx.config.project.project_code})")
        doc.add_paragraph(f"Run: {ctx.run_id}")
        doc.add_paragraph(f"Target CRS: {ctx.config.crs.target_authority}")
        if dry_run:
            doc.add_paragraph("Dry run: no campaign source was opened.")
        else:
            assert self._survey is not None
            doc.add_heading("Acquired campaign", level=2)
            values = (
                ("Processed blocks", len(self._survey.blocks)),
                (
                    "Blocks with point clouds",
                    sum(b.point_cloud is not None for b in self._survey.blocks),
                ),
                ("L2 blocks", sum(b.sensor == "L2" for b in self._survey.blocks)),
                ("L3 blocks", sum(b.sensor == "L3" for b in self._survey.blocks)),
                ("Control files", len(self._survey.control_files)),
                ("Daily report files", len(self._survey.daily_report_files)),
                ("Phase 04 prospect polygons", self._prospect_count),
            )
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            table.rows[0].cells[0].text = "Item"
            table.rows[0].cells[1].text = "Count"
            for label, value in values:
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = str(value)
            doc.add_heading("Produced in this slice", level=2)
            for item in (
                "Block asset and data-gap register",
                "Spatial block coverage GeoPackage",
                "Cloud-Optimized DTM, hillshade and slope",
                "Structure/outcrop interpretation and field-planning GeoPackage schema",
                "Acquired coverage / flight-record PDF",
                "QGIS project with a named DTM-backed 3D terrain view",
            ):
                doc.add_paragraph(item, style="List Bullet")
            doc.add_heading("Remaining functional products", level=2)
            for item in (
                "Build a full-area orthomosaic from selected per-block orthos.",
                "Build a full-area DSM mosaic.",
                "Convert per-block LAS to LAZ/COPC when PDAL is installed.",
                "Populate geological structure/outcrop interpretations from imagery and terrain.",
            ):
                doc.add_paragraph(item, style="List Bullet")
        doc.save(str(path))
        return path


PHASE = Phase05DroneLidar
